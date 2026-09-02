#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
make_fixtures.py · 生成 industry-report-update-skill 评估 fixtures

用法：
    python evals/cases/make_fixtures.py

生成：
    evals/cases/fixtures_positive/        # 正向：合规 + 数值一致 + 格式保真
        ├── old.docx
        ├── new.docx
        ├── sources/
        │   ├── 章节-上交所.csv
        │   └── 章节-深交所.csv
        └── 溯源.jsonl

    evals/cases/fixtures_negative/        # 负向：幽灵条目 / 金额矛盾
        ├── ghost_entity.docx
        ├── sources/
        │   └── 章节-上交所.csv
        └── 溯源_ghost.jsonl

        value_mismatch.docx
        sources/
            └── 章节-上交所.csv
        溯源_mismatch.jsonl

防回归：
    - 正向 5/5 必须全部通过 11 道门禁
    - 负向 2/2 必须被对应门禁拦截
"""
import json
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).parent
FIXTURES_POS = ROOT / "fixtures_positive"
FIXTURES_NEG = ROOT / "fixtures_negative"


def make_old_docx(path: Path) -> None:
    """生成合规的旧月报（含 5 行 IPO 审核 + 一致性数值 + 来源脚注）"""
    doc = Document()

    # 标题
    title = doc.add_heading("行业月报（合规示例）", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 摘要段
    doc.add_paragraph("本期共有 5 家公司在审。数据来源：示例交易所，截至 2026-08-31。")

    # 表格
    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"
    headers = ["公司简称", "板块", "审核状态", "更新日期"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    rows_data = [
        ("A 公司", "科创板", "已问询", "2026-08-20"),
        ("B 公司", "创业板", "已受理", "2026-08-22"),
        ("C 公司", "科创板", "上市委会议", "2026-08-25"),
        ("D 公司", "北交所", "已问询", "2026-08-18"),
        ("E 公司", "创业板", "提交注册", "2026-08-28"),
    ]
    for r_idx, row in enumerate(rows_data, start=1):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = value
            # 显式设字号（保格式改写必备）
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10.5)

    # 数据来源
    p = doc.add_paragraph()
    p.add_run("数据来源：示例交易所，截至 2026-08-31。").italic = True

    doc.save(str(path))


def make_new_docx(path: Path, ghost: bool = False, mismatch: bool = False) -> None:
    """生成新月报
    ghost=True: 加 1 行幽灵条目（共 6 家，但正文说"共 5 家"）—— 交叉一致性应拦截
    mismatch=True: 改 1 行金额与溯源.jsonl 不一致 —— verify_value 应拦截
    """
    doc = Document()
    title = doc.add_heading("行业月报（合规示例）", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 摘要段
    doc.add_paragraph("本期共有 5 家公司在审。数据来源：示例交易所，截至 2026-09-30。")

    # 表格（6 行 if ghost）
    n = 6 if ghost else 5
    table = doc.add_table(rows=n + 1, cols=4)
    table.style = "Table Grid"
    headers = ["公司简称", "板块", "审核状态", "更新日期"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    rows_data = [
        ("A 公司", "科创板", "已问询", "2026-09-20"),
        ("B 公司", "创业板", "已受理", "2026-09-22"),
        ("C 公司", "科创板", "上市委会议", "2026-09-25"),
        ("D 公司", "北交所", "已问询", "2026-09-18"),
        ("E 公司", "创业板", "注册生效", "2026-09-28"),
    ]
    if ghost:
        rows_data.append(("F 公司", "科创板", "已问询", "2026-09-15"))  # 幽灵

    for r_idx, row in enumerate(rows_data, start=1):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = value
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.add_run("数据来源：示例交易所，截至 2026-09-30。").italic = True
    doc.save(str(path))


def make_csv(path: Path, ghost: bool = False, mismatch: bool = False) -> None:
    """生成源 CSV（与新月报表格对应）"""
    rows = [
        "公司简称,板块,审核状态,更新日期",
        "A 公司,科创板,已问询,2026-09-20",
        "B 公司,创业板,已受理,2026-09-22",
        "C 公司,科创板,上市委会议,2026-09-25",
        "D 公司,北交所,已问询,2026-09-18",
        "E 公司,创业板,注册生效,2026-09-28",
    ]
    if ghost:
        # ghost 的 CSV 里没有 F 公司（因为是幽灵条目）
        pass
    if mismatch:
        # mismatch 的 CSV：E 公司金额故意改 1 个字段
        for i, line in enumerate(rows):
            if line.startswith("E 公司"):
                rows[i] = "E 公司,创业板,注册生效,2026-09-30"  # 改日期
    path.write_text("\n".join(rows) + "\n", encoding="utf-8")


def make_traceability_jsonl(path: Path, mismatch: bool = False) -> None:
    """生成溯源.jsonl（带锚点）"""
    rows_data = [
        ("A 公司", "已问询", "2026-09-20"),
        ("B 公司", "已受理", "2026-09-22"),
        ("C 公司", "上市委会议", "2026-09-25"),
        ("D 公司", "已问询", "2026-09-18"),
        ("E 公司", "注册生效", "2026-09-28"),
    ]
    records = []
    for company, status, date in rows_data:
        records.append({
            "cell": f"表1!C{rows_data.index((company, status, date)) + 2}",
            "value": status,
            "source_file": "sources/章节-上交所.csv",
            "source_key": company,
            "source_field": "审核状态",
            "status": "verified",
            "cross_checked": ["示例交易所"],
            "as_of": "2026-09-30",
        })
        if mismatch and company == "E 公司":
            # mismatch: 改 date 字段让其对不上
            records[-1]["value"] = "注册生效"
            records[-1]["source_field"] = "审核状态"
            # 但 JSONL 里写一个错的 as_of
            records[-1]["as_of"] = "2026-09-29"
    path.write_text(
        "\n".join(json.dumps(r, ensure_ascii=False) for r in records) + "\n",
        encoding="utf-8",
    )


def setup_positive():
    if FIXTURES_POS.exists():
        shutil.rmtree(FIXTURES_POS)
    FIXTURES_POS.mkdir(parents=True)
    (FIXTURES_POS / "sources").mkdir()

    old_path = FIXTURES_POS / "old.docx"
    new_path = FIXTURES_POS / "new.docx"
    csv_path = FIXTURES_POS / "sources" / "章节-上交所.csv"
    jsonl_path = FIXTURES_POS / "溯源.jsonl"
    csv_path2 = FIXTURES_POS / "sources" / "章节-深交所.csv"

    make_old_docx(old_path)
    make_new_docx(new_path)
    make_csv(csv_path)
    csv_path2.write_text("公司简称,板块,审核状态,更新日期\n", encoding="utf-8")
    make_traceability_jsonl(jsonl_path)

    print(f"[POS] {old_path}")
    print(f"[POS] {new_path}")
    print(f"[POS] {csv_path}")
    print(f"[POS] {csv_path2}")
    print(f"[POS] {jsonl_path}")


def setup_negative_ghost():
    neg_dir = FIXTURES_NEG / "ghost"
    if neg_dir.exists():
        shutil.rmtree(neg_dir)
    neg_dir.mkdir(parents=True)
    (neg_dir / "sources").mkdir()

    new_path = neg_dir / "ghost_entity.docx"
    csv_path = neg_dir / "sources" / "章节-上交所.csv"
    jsonl_path = neg_dir / "溯源_ghost.jsonl"

    make_new_docx(new_path, ghost=True)
    make_csv(csv_path, ghost=True)  # CSV 里没有 F 公司（幽灵）
    make_traceability_jsonl(jsonl_path)

    print(f"[NEG-ghost] {new_path}")
    print(f"[NEG-ghost] {csv_path}")
    print(f"[NEG-ghost] {jsonl_path}")


def setup_negative_mismatch():
    neg_dir = FIXTURES_NEG / "mismatch"
    if neg_dir.exists():
        shutil.rmtree(neg_dir)
    neg_dir.mkdir(parents=True)
    (neg_dir / "sources").mkdir()

    new_path = neg_dir / "value_mismatch.docx"
    csv_path = neg_dir / "sources" / "章节-上交所.csv"
    jsonl_path = neg_dir / "溯源_mismatch.jsonl"

    make_new_docx(new_path, mismatch=True)
    make_csv(csv_path, mismatch=True)
    make_traceability_jsonl(jsonl_path, mismatch=True)

    print(f"[NEG-mismatch] {new_path}")
    print(f"[NEG-mismatch] {csv_path}")
    print(f"[NEG-mismatch] {jsonl_path}")


def main():
    print("Generating industry-report-update-skill evaluation fixtures...")
    setup_positive()
    setup_negative_ghost()
    setup_negative_mismatch()
    print("Done.")


if __name__ == "__main__":
    main()