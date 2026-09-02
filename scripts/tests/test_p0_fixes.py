# -*- coding: utf-8 -*-
"""
test_p0_fixes.py —— P0 修复回归测试（python scripts/tests/test_p0_fixes.py）
覆盖：P0-1 archive 守卫/报告通配；P0-3 docx_utils 超链接 run 改写；P0-4 verify_value 单位识别；
      P0-5 yaml_lite.mini == pyyaml（含 tool_registry.yaml 冒号值）。
"""
import io
import os
import sys

HERE = os.path.dirname(os.path.abspath(__file__))
SCRIPTS = os.path.dirname(HERE)
BASE = os.path.dirname(SCRIPTS)
sys.path.insert(0, SCRIPTS)

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

failures = []


def check(name, cond, detail=''):
    if not cond:
        failures.append('%s: %s' % (name, detail))
        print('  ❌ %s %s' % (name, detail))
    else:
        print('  ✅ %s' % name)


# ---------------- P0-1：archive_to_workspace ----------------
def test_p01():
    print('\n== P0-1 archive_to_workspace ==')
    import archive_to_workspace as A
    check('_md_is_report 通道健康度通配',
          A._md_is_report('通道健康度-2026-08.md') is True)
    check('_md_is_report 精确名', A._md_is_report('格式对比报告.md') is True)
    check('_md_is_report 排除', A._md_is_report('foo.md') is False)
    inside = os.path.join(A.BASE, 'runs', 'x')
    check('_inside_base 子目录(反斜杠)', A._inside_base(inside) is True)
    check('_inside_base 等于BASE', A._inside_base(A.BASE) is True)
    outside = r'D:\Desktop\industry-research\reports\old-monthly-report-2026-08'
    check('_inside_base 非基目录', A._inside_base(outside) is False)
    # 混合分隔符（Windows 混淆路径）
    mixed = os.path.join(A.BASE.replace('\\', '/'), 'runs', 'x')
    check('_inside_base 混合分隔符', A._inside_base(mixed) is True)
    # 兄弟目录共享前缀（如 ...industry-report-update-x）必须判为外部（段边界）
    sibling = os.path.join(os.path.dirname(A.BASE), 'industry-report-update-x')
    check('_inside_base 兄弟前缀目录(应False)', A._inside_base(sibling) is False, 'sibling=%s' % sibling)


# ---------------- P0-3：docx_utils 超链接 run 改写 ----------------
def _add_hyperlink(p, text, url='https://example.com'):
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('r:id'), 'rId9')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    hl.append(r)
    p._element.append(hl)


def _para_text(p):
    return ''.join((t.text or '') for r in p._element.iter(qn('w:r')) for t in r.iter(qn('w:t')))


def test_p03():
    print('\n== P0-3 docx_utils hyperlink ==')
    from docx_utils import set_para_text_keep_fmt, set_cell_text_keep_fmt
    d = Document()
    p = d.add_paragraph()
    p.add_run('前导')
    _add_hyperlink(p, '旧链接文字')
    set_para_text_keep_fmt(p, '新正文')
    texts = _para_text(p)
    check('段落改写后超链接run旧文字被清理', '旧链接文字' not in texts)
    check('段落改写后包含新正文', '新正文' in texts)

    # 单元格
    d2 = Document()
    table = d2.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cp = cell.paragraphs[0]
    cp.add_run('前导')
    _add_hyperlink(cp, '旧链接文字')
    set_cell_text_keep_fmt(cell, '新值')
    ctexts = _para_text(cp)
    check('单元格改写后超链接run旧文字被清理', '旧链接文字' not in ctexts)
    check('单元格改写后包含新值', '新值' in ctexts)


# ---------------- P0-4：verify_value 单位识别 ----------------
def test_p04():
    print('\n== P0-4 verify_value unit ==')
    from audit import verify_value as V
    # _unit_factor 应正确区分 亿美元(币种) 与 亿(数量级)
    check('亿美元 识别且factor=1e8',
          V._unit_factor('1.2亿美元')[0] == 1e8)
    check('亿元 识别', V._unit_factor('3.3亿元')[0] == 1e8)
    check('万 识别', V._unit_factor('5万')[0] == 1e4)
    check('元 识别', V._unit_factor('150.80元')[0] == 1.0)
    # 币种区分：等值异币必须判"不一致"(硬伤)
    ok, note, hard = V._compare('1.2亿美元', '1.2亿元')
    check('等值异币(亿美元vs亿元)判不一致', (ok is False) and (hard is True), 'ok=%s note=%s' % (ok, note))
    # 同币种同量级应一致
    ok2, note2, hard2 = V._compare('1.2亿元', '1.2亿元')
    check('同币同量级判一致', ok2 is True and hard2 is False, note2)


# ---------------- P0-5：yaml_lite.mini == pyyaml ----------------
def test_p05():
    print('\n== P0-5 yaml_lite.mini == pyyaml ==')
    import yaml_lite
    import yaml as pyyaml
    # 强制 mini 后端（保存/恢复），否则 PyYAML 在场时 mini==pyyaml 实为 pyyaml vs pyyaml
    saved_pyyaml = getattr(yaml_lite, '_pyyaml', None)
    yaml_lite._pyyaml = None
    diffs = []
    scan_dirs = ['config', 'config/口径快照', 'packs/robotics']
    for sub in scan_dirs:
        d = os.path.join(BASE, sub)
        if not os.path.isdir(d):
            continue
        for f in sorted(os.listdir(d)):
            if not f.endswith('.yaml'):
                continue
            path = os.path.join(d, f)
            try:
                mini = yaml_lite.load_yaml(path)
            except Exception as e:
                diffs.append('%s/%s: mini 解析失败 %s' % (sub, f, e))
                continue
            try:
                py = pyyaml.safe_load(open(path, encoding='utf-8'))
            except Exception as e:
                diffs.append('%s/%s: pyyaml 解析失败 %s' % (sub, f, e))
                continue
            if mini != py:
                diffs.append('%s/%s: mini != pyyaml' % (sub, f))
    yaml_lite._pyyaml = saved_pyyaml
    check('全部 config/口径快照/packs yaml 的 mini==pyyaml', not diffs, '; '.join(diffs[:6]))


# ---------------- P1-1：build_traceability 数值列锚点 → verify_value 回读 ----------------
def test_p11():
    print('\n== P1-1 build_traceability 数值锚点 + verify_value 回读 ==')
    import subprocess, tempfile, json, csv as csvm
    tmp = tempfile.mkdtemp()
    # ① 构造新月报 docx（表1：公司简称/营业收入/净利润）
    d = Document()
    t = d.add_table(rows=3, cols=3)
    for i, h in enumerate(['公司简称', '营业收入', '净利润']):
        t.cell(0, i).text = h
    t.cell(1, 0).text = 'A 公司'; t.cell(1, 1).text = '16.99亿元'; t.cell(1, 2).text = '5.91亿元'
    t.cell(2, 0).text = 'B 公司'; t.cell(2, 1).text = '10.00亿元'; t.cell(2, 2).text = '-1.00亿元'
    docx_path = os.path.join(tmp, '新.docx')
    d.save(docx_path)
    # ② 构造源文件 csv（键列=公司简称，含 营业收入/净利润 列）
    csv_path = os.path.join(tmp, '源.csv')
    with open(csv_path, 'w', encoding='utf-8-sig', newline='') as f:
        w = csvm.writer(f)
        w.writerow(['公司简称', '营业收入', '净利润'])
        w.writerow(['A 公司', '16.99亿元', '5.91亿元'])
        w.writerow(['B 公司', '10.00亿元', '-1.00亿元'])
    # ③ 映射（value_cols=营业收入/净利润）
    map_path = os.path.join(tmp, 'm.json')
    json.dump({"1": {"source_file": '源.csv', "key_col": '公司简称', "value_cols": ['营业收入', '净利润'],
                     "unit": '亿元', "reporting_period": '2025年报', "as_of": '2026-08-31'}}, open(map_path, 'w', encoding='utf-8'))
    jsonl = os.path.join(tmp, '溯源.jsonl')
    subprocess.run([sys.executable, os.path.join(SCRIPTS, 'audit', 'build_traceability.py'), docx_path,
                    '--mapping', map_path, '--base-dir', tmp, '--out', jsonl], check=True)
    recs = [json.loads(l) for l in open(jsonl, encoding='utf-8') if l.strip()]
    # 数值列应生成带 value+source_field+source_key 的记录
    num_recs = [r for r in recs if r.get('source_field') in ('营业收入', '净利润')]
    check('数值列记录含 source_key', all(r.get('source_key') for r in num_recs))
    check('数值列记录含真实value+source_field',
          any(r.get('source_field') == '营业收入' and '16.99' in (r.get('value') or '') for r in num_recs))
    # ④ verify_value 回读：应全部一致（ok），数值回读闭环
    v = subprocess.run([sys.executable, os.path.join(SCRIPTS, 'audit', 'verify_value.py'), jsonl,
                        '--base-dir', tmp], capture_output=True, text=True)
    out = (v.stdout or '') + (v.stderr or '')
    check('verify_value 数值回读无硬伤', '❌ 不一致或锚点失败 0' in out and '✅ 回读一致' in out, out.strip()[:160])


if __name__ == '__main__':
    test_p01()
    test_p03()
    test_p04()
    test_p05()
    test_p11()
    print('\n========= 结果：%s =========' % ('PASS' if not failures else 'FAIL %d 项' % len(failures)))
    sys.exit(1 if failures else 0)
