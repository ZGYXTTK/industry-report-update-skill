# -*- coding: utf-8 -*-
"""
docx_build.py —— 月报构建原语库（P2-1）

把 2026-08 期 build_aug.py 中验证过的保格式操作抽象为可复用原语，
供 build_report.py（声明式引擎）与行业包自定义钩子共用。

设计原则（全部来自 2026-08 期实测教训）：
  1. 段落写入必须 run 保留（set_para_keep）——单 segment 重建使 format_diff 跌至 94.1%；
  2. 表格写入必须保留段落/run 结构（set_cell_keep）——旧值残留使一致性门禁误报；
  3. 表格重填前 strip_vmerge + fill_table（删残行/克隆补行）——索引错配与残行根因；
  4. 残句删除用包含匹配（in），不用前缀匹配；
  5. 表格定位支持「序号」与「表头特征」双模式（防结构性编辑后序号漂移）。
"""
import copy
import os

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

import docx_utils as du


class DocEditor:
    """月报 docx 的保格式编辑器：包装 Document + 全部验证过的构建原语。"""

    def __init__(self, template_path):
        self.template_path = template_path
        self.doc = Document(template_path)

    # ---------- 段落 ----------
    def paras(self):
        return list(self.doc.paragraphs)

    def find_para(self, prefix, start=0):
        for i, p in enumerate(self.paras()):
            if i >= start and p.text.strip().startswith(prefix):
                return i
        raise ValueError('段落未找到（前缀）：%s' % prefix[:40])

    def find_para_contains(self, sub, start=0):
        for i, p in enumerate(self.paras()):
            if i >= start and sub in p.text:
                return i
        raise ValueError('段落未找到（包含）：%s' % sub[:40])

    @staticmethod
    def set_text(p, text):
        """run 保留式段落写入。"""
        du.set_para_text_keep_fmt(p, text)

    def insert_after(self, ref_p, text):
        el = copy.deepcopy(ref_p._p)
        ref_p._p.addnext(el)
        np = Paragraph(el, ref_p._parent)
        self.set_text(np, text)
        return np

    def insert_before(self, ref_p, text):
        el = copy.deepcopy(ref_p._p)
        ref_p._p.addprevious(el)
        np = Paragraph(el, ref_p._parent)
        self.set_text(np, text)
        return np

    def delete_para(self, p):
        p._p.getparent().remove(p._p)

    def replace_range(self, start_prefix, end_prefix, items):
        """把 start 锚点之后到 end 锚点之前的全部段落替换为 items（继承 start 锚点格式）。"""
        ps = self.paras()
        si = self.find_para(start_prefix)
        ei = si + 1 + next((k for k, p in enumerate(ps[si + 1:]) if p is ps[ei] if False), 0) if False else None
        # 直接定位
        start_p = ps[si]
        ei = None
        for k in range(si + 1, len(ps)):
            if ps[k].text.strip().startswith(end_prefix):
                ei = k
                break
        if ei is None:
            raise ValueError('区间终点未找到：%s' % end_prefix[:40])
        end_p = ps[ei]
        for p in ps[si + 1:ei]:
            self.delete_para(p)
        cur = start_p
        for it in items:
            if isinstance(it, dict):
                text = it.get("text", "")
                overrides = it.get("fmt")
            else:
                text = it
                overrides = None
            cur = self.insert_after(cur, text)
            if overrides and cur.runs:
                run0 = cur.runs[0]
                rfont_val = None
                if isinstance(overrides, dict):
                    rfont_val = overrides.pop("rFonts", None) if "rFonts" in overrides else None
                du._apply_rpr_overrides(run0, overrides or {})
                if rfont_val:
                    rpr_el = run0._element.find(qn("w:rPr"))
                    if rpr_el is None:
                        rpr_el = run0._element.makeelement(qn("w:rPr"), {})
                        run0._element.insert(0, rpr_el)
                    rf = rpr_el.find(qn("w:rFonts"))
                    if rf is None:
                        rf = rpr_el.makeelement(qn("w:rFonts"), {})
                        rpr_el.append(rf)
                    rf.set(qn("w:ascii"), str(rfont_val))
                    rf.set(qn("w:eastAsia"), str(rfont_val))
        return cur

    def rep_contains(self, sub, text):
        for p in self.paras():
            if sub in p.text:
                self.set_text(p, text)
                return True
        print('⚠️  rep_contains 未命中：%s' % sub[:30])
        return False

    def rep_prefix(self, prefix, text):
        for p in self.paras():
            if p.text.strip().startswith(prefix):
                self.set_text(p, text)
                return True
        print('⚠️  rep_prefix 未命中：%s' % prefix[:30])
        return False

    def delete_contains(self, patterns, start_prefix=None):
        """删除包含任一 pattern 的段落；start_prefix 限定起始锚点（其后才删）。"""
        ps = self.paras()
        begin = 0
        if start_prefix:
            begin = self.find_para(start_prefix) + 1
        killed = 0
        for p in self.paras()[begin:]:
            t = p.text.strip()
            if any(k in t for k in patterns):
                self.delete_para(p)
                killed += 1
        return killed

    # ---------- 表格 ----------
    def find_table(self, index=None, header_contains=None):
        """按序号或表头特征定位表格（header_contains 匹配首行任意单元格）。"""
        tables = self.doc.tables
        if index is not None:
            return tables[index]
        if header_contains:
            for t in tables:
                if t.rows and any(header_contains in c.text for c in t.rows[0].cells):
                    return t
        raise ValueError('表格未定位（index=%r header_contains=%r）' % (index, header_contains))

    @staticmethod
    def strip_vmerge(table):
        return du.strip_vmerge(table)

    @staticmethod
    def fill_table(table, rows, start_row=1):
        return du.fill_table(table, rows, start_row=start_row)

    @staticmethod
    def clone_row(table):
        tr = table.rows[-1]._tr
        table.rows[-1]._tr.addnext(copy.deepcopy(tr))

    @staticmethod
    def del_row(table, r):
        tr = table.rows[r]._tr
        tr.getparent().remove(tr)

    def fill_table_safe(self, table, rows, start_row=1):
        """fill_table + 行数自洽（等价于 build 期的 while 删行逻辑，见 fill_table）。"""
        return self.fill_table(table, rows, start_row=start_row)

    def insert_note_after_table(self, table, text, tmpl_p=None):
        """在表格之后插入说明段（继承 Body Text 模板段格式）。"""
        if tmpl_p is None:
            for p in self.paras():
                if p.style is not None and p.style.name and 'Body Text' in p.style.name and p.text.strip():
                    tmpl_p = p
                    break
        el = copy.deepcopy(tmpl_p._p)
        table._tbl.addnext(el)
        np = Paragraph(el, tmpl_p._parent)
        self.set_text(np, text)
        return np

    # ---------- 图片 ----------
    def replace_image_part(self, target_name, image_path):
        """替换 word/media/<name>（JPEG/PNG 字节级替换；尺寸由 Word 框架约束）。"""
        from PIL import Image
        ext = os.path.splitext(target_name)[1].lower()
        img = Image.open(image_path)
        if ext in ('.jpeg', '.jpg') and img.format != 'JPEG':
            buf = image_path + '.conv.jpg'
            img.convert('RGB').save(buf, 'JPEG', quality=92)
            image_path = buf
        data = open(image_path, 'rb').read()
        for rel in self.doc.part.rels.values():
            if rel.reltype.endswith('/image') and rel.target_ref.endswith(os.path.basename(target_name)):
                rel.target_part._blob = data
                return True
        print('⚠️  图片 part 未找到：%s' % target_name)
        return False

    # ---------- 输出 ----------
    def save(self, out_path):
        os.makedirs(os.path.dirname(out_path), exist_ok=True)
        self.doc.save(out_path)
        return out_path
