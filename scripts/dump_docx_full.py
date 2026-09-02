# -*- coding: utf-8 -*-
"""Full extraction: all tables (every cell), all paragraphs (full text), charts detection."""
import sys, io, zipfile, re
from docx import Document

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

path = sys.argv[1]
doc = Document(path)

print("=" * 30, "PARAGRAPHS (full)", "=" * 30)
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if t.strip():
        print(f"[P{i:04d}] {t}")

print("=" * 30, "TABLES (full)", "=" * 30)
for ti, tbl in enumerate(doc.tables):
    print(f"\n--- T{ti} ({len(tbl.rows)}x{len(tbl.columns)}) ---")
    for ri, row in enumerate(tbl.rows):
        cells = []
        for c in row.cells:
            txt = c.text.replace("\n", "␤").strip()
            cells.append(txt)
        # dedupe merged repeats
        dedup = []
        prev = object()
        for c in cells:
            if c != prev:
                dedup.append(c)
            prev = c
        print(f"R{ri}: " + " ║ ".join(dedup))

print("=" * 30, "CHARTS / IMAGES", "=" * 30)
z = zipfile.ZipFile(path)
names = z.namelist()
charts = [n for n in names if 'chart' in n.lower()]
media = [n for n in names if n.startswith('word/media/')]
print("chart parts:", charts)
print("media parts:", media)
