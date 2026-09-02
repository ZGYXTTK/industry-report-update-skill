# -*- coding: utf-8 -*-
"""
extract_numbers.py —— 数据流审计第①件：结论性语句提取（v2：召回提升 + 结构化输出）

v2 改进：
  1. 补「状态词在前」语序（如「在审企业共16家」），不再只认「16家在审」；
  2. 去掉「每单元只取首个命中」，改为全量匹配并按位置去重排序；
  3. 标记「合计/总计」行；
  4. 除 .md 外，额外输出 结论.jsonl（{loc,kind,fragment,ctx,is_total}），供程序化填充审计模板。

用法:
    python extract_numbers.py 旧月报.docx [--out 结论性语句清单.md] [--json 结论.jsonl]
"""
import argparse
import json
import re

from docx import Document

# 长词在前，避免「辅导备案」被「辅导」截断
STATE_WORDS = '辅导备案|注册生效|再融资|在审|辅导|备案|发行|申报|上市|终止|过会|受理|并购|问询'
TOTAL_KEYS = ('合计', '总计', '小计')

PATTERNS = [
    # 家数-状态后置：16家在审 / 16家公司辅导
    (r'\d+(?:\.\d+)?\s*家(?:公司|企业)?(?:' + STATE_WORDS + r')', '家数-状态'),
    # 家数-状态前置：在审企业共16家 / 辅导备案20家
    (r'(?:' + STATE_WORDS + r')(?:企业|公司)?(?:共|合计|达|约)?\s*\d+(?:\.\d+)?\s*家', '家数-状态前置'),
    # 家数-截至：截至6月末共16家
    (r'截至[^。，；]{0,20}?共\s*\d+(?:\.\d+)?\s*家', '家数-截至'),
    # 家数-共：共16家
    (r'共\s*\d+(?:\.\d+)?\s*家', '家数-共'),
    # 笔数：12笔
    (r'\d+(?:\.\d+)?\s*笔', '笔数'),
    # 金额：1.2亿元 / 3000万美元
    (r'\d+(?:\.\d+)?\s*(?:亿元|万元|亿美元|亿港元|港元|万美元)', '金额'),
    # 同比环比：同比增长12.3% / 环比-2.1%
    (r'(?:环比|同比)[^\d]{0,8}[+-]?\d+(?:\.\d+)?%?', '同比环比'),
    # 百分比：12.3%
    (r'\d+(?:\.\d+)?\s*%', '百分比'),
]


def _col_letter(ci):
    s = ''
    n = ci + 1
    while n > 0:
        n, r = divmod(n - 1, 26)
        s = chr(65 + r) + s
    return s


def _extract(text):
    """全量匹配、按位置去重排序。返回 [(kind, fragment, start, end), ...]。"""
    seen = set()
    results = []
    for pat, kind in PATTERNS:
        for m in re.finditer(pat, text):
            span = (m.start(), m.end())
            if span in seen:
                continue
            seen.add(span)
            results.append((kind, m.group(0), m.start(), m.end()))
    results.sort(key=lambda x: x[2])
    return results


def _is_total(text):
    return any(k in text for k in TOTAL_KEYS)


def iter_hits(doc):
    """产出 (loc, text, is_total_row)。"""
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            # 段落不是表格行，不标记合计（避免「合计16家」这类短语误标）
            yield ('段落%d' % i, p.text, False)
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            row_is_total = any(_is_total(c.text) for c in row.cells)
            for ci, cell in enumerate(row.cells):
                t = cell.text.strip()
                if t:
                    yield ('表%d!%s%d' % (ti + 1, _col_letter(ci), ri + 1), t, row_is_total)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('docx')
    ap.add_argument('--out', default='结论性语句清单.md')
    ap.add_argument('--json', default='结论.jsonl')
    args = ap.parse_args()

    doc = Document(args.docx)
    hits = []
    for loc, text, is_total in iter_hits(doc):
        for kind, frag, s, e in _extract(text):
            ctx = text[max(0, s - 20): e + 20]
            hits.append({'loc': loc, 'kind': kind, 'fragment': frag,
                         'ctx': ctx, 'is_total': is_total})

    lines = [
        '# 结论性语句清单（自动提取，请补全「数据源/口径/时点」三列）',
        '',
        '| # | 位置 | 类别 | 命中片段 | 上下文 | 合计行 | 数据源 | 口径 | 时点型? |',
        '| --- | --- | --- | --- | --- | --- | --- | --- | --- |',
    ]
    for i, h in enumerate(hits, 1):
        lines.append('| %d | %s | %s | %s | %s | %s |  |  | ✅/❌ |'
                     % (i, h['loc'], h['kind'], h['fragment'].replace('|', '\\|'),
                        h['ctx'].replace('|', '\\|'), '✓' if h['is_total'] else ''))
    lines += ['', '> 命中 %d 条。请按《templates/数据流审计产出模板.md》补全右侧三列。' % len(hits)]

    with open(args.out, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    with open(args.json, 'w', encoding='utf-8') as f:
        for h in hits:
            f.write(json.dumps(h, ensure_ascii=False) + '\n')
    print('已写出 %s 与 %s，共命中 %d 条' % (args.out, args.json, len(hits)))


if __name__ == '__main__':
    main()
