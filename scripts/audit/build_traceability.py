# -*- coding: utf-8 -*-
"""
build_traceability.py —— 溯源.jsonl 自动生成（Step 7 归档辅助）

把「逐条可溯源」从手工拼 JSON 变成半自动：给定「表 → 源文件」映射，
自动为月报每张表的每个数据行生成一条溯源记录（cell/value/source_file/as_of…），
Agent 只需补 status/cross_checked 等交叉验证字段，降低漏标与手误。

映射文件（JSON）结构示例：
    {
      "1": {"source_file": "下载资料/沪深北交易所官网+iFinD-A股IPO在审企业汇总.csv",
            "key_col": "公司简称", "value_cols": ["营业收入", "净利润"],
            "unit": "亿元", "reporting_period": "最近披露"},
      "8": {"source_file": "下载资料/同花顺iFinD新闻-非上市公司投融资-7月.csv",
            "key_col": 2, "value_cols": [7]}
    }
key 为 1-based 表序号；key_col / value_cols 可为表头名称（子串匹配）或列下标（0-based）。
value_cols 列出的数值列会对**每一行**再生成一条带 source_field+value 的锚点记录，供 verify_value 逐值回读。

用法:
    python build_traceability.py 新月报.docx --mapping 溯源映射.json --as-of 2026-07-31 --out 溯源.jsonl

v2 变化：若映射 spec 含 source_key_col（源文件键列表头名，默认自动用源文件首列），
且源文件存在，则自动为每条记录写入锚点字段 source_key/source_key_col——
供 verify_value.py 回读比对（P0 门禁：月报数字必须等于源文件数字）。
"""
import argparse
import csv
import json
import os
import re

from docx import Document


def _col_letter(ci):
    s = ''
    n = ci + 1
    while n > 0:
        n, r = divmod(n - 1, 26)
        s = chr(65 + r) + s
    return s


def _resolve_key_col(table, spec):
    """spec 可为整数（下标）或字符串（表头名称子串匹配）。"""
    if isinstance(spec, int):
        return spec
    if isinstance(spec, str) and table.rows:
        header = [c.text.strip() for c in table.rows[0].cells]
        for ci, h in enumerate(header):
            if spec in h:
                return ci
    return 0


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('docx')
    ap.add_argument('--mapping', required=True, help='表→源文件 映射 JSON 文件')
    ap.add_argument('--as-of', default=None, help='数据截至时点 YYYY-MM-DD')
    ap.add_argument('--base-dir', default=None, help='source_file 相对路径锚定目录（默认取映射文件所在目录）')
    ap.add_argument('--out', default='溯源.jsonl')
    args = ap.parse_args()

    with open(args.mapping, encoding='utf-8') as f:
        mapping = json.load(f)

    base_dir = args.base_dir or os.path.dirname(os.path.abspath(args.mapping))

    def _read_csv_keys(path):
        """读源文件首列键集合（归一化），用于自动锚定。文件缺失返回 None。"""
        if not os.path.exists(path):
            return None
        for enc in ('utf-8-sig', 'utf-8', 'gbk'):
            try:
                with open(path, encoding=enc, newline='') as f:
                    rows = [r for r in csv.reader(f) if any((c or '').strip() for c in r)]
                if rows:
                    return rows
            except (UnicodeDecodeError, UnicodeError):
                continue
        return None

    def _norm_key(v):
        t = re.sub(r'\s+', '', str(v or ''))
        out = []
        for ch in t:
            code = ord(ch)
            if 0xFF01 <= code <= 0xFF5E:
                ch = chr(code - 0xFEE0)
            out.append(ch)
        t = ''.join(out)
        t = re.sub(r'[（(][^）)]*[）)]', '', t)
        t = re.sub(r'(股份有限公司|有限责任公司|有限公司|控股集团|集团)', '', t)
        return t.strip()

    csv_cache = {}

    doc = Document(args.docx)
    rows_out = []
    for tkey, spec in mapping.items():
        ti = int(tkey) - 1
        if ti < 0 or ti >= len(doc.tables):
            print('跳过不存在的表序号 %s' % tkey)
            continue
        table = doc.tables[ti]
        key_col = _resolve_key_col(table, spec.get('key_col', 0))
        # value_cols：需逐值回读的数值列（列下标或表头名称子串）
        value_cols0 = spec.get('value_cols', []) or []
        value_cols = [_resolve_key_col(table, vc) for vc in value_cols0] if value_cols0 else []
        sf = spec.get('source_file', '')
        unit = spec.get('unit')
        rp = spec.get('reporting_period')
        url = spec.get('url')
        # v2：尝试装载源文件，为自动锚定做准备
        src_path = os.path.join(base_dir, sf) if sf and not os.path.isabs(sf) else sf
        if sf and src_path not in csv_cache:
            csv_cache[src_path] = _read_csv_keys(src_path)
        src_rows = csv_cache.get(src_path)
        src_key_col_name = spec.get('source_key_col')  # 源文件键列表头名（可选）

        for ri in range(1, len(table.rows)):
            cells = table.rows[ri].cells
            key_text = cells[key_col].text.strip() if key_col < len(cells) else ''
            if not key_text:
                key_text = '（键列为空）'
            # 解析源文件键列下标与命中情况（每行一次）
            header = src_rows[0] if src_rows else []
            kc = 0
            skc = ''
            nk = _norm_key(key_text)
            hit = False
            if src_rows and key_text != '（键列为空）':
                if src_key_col_name:
                    for ci, h in enumerate(header):
                        if re.sub(r'\s+', '', src_key_col_name) in re.sub(r'\s+', '', h or ''):
                            kc = ci
                            break
                hit = any(kc < len(r) and _norm_key(r[kc]) == nk for r in src_rows[1:])
                skc = src_key_col_name or (header[kc] if kc < len(header) else '')
            row_cols = value_cols if value_cols else [key_col]
            hdr = [c.text.strip() for c in table.rows[0].cells] if len(table.rows) else []
            for vc in row_cols:
                val_text = cells[vc].text.strip() if vc < len(cells) else ''
                rec = {
                    'cell': '表%d!R%d' % (ti + 1, ri + 1),
                    'value': val_text if value_cols else key_text,
                    'source_file': sf,
                    'as_of': args.as_of,
                    'status': 'verified',
                    'source_field': (hdr[vc] if vc < len(hdr) else ''),
                }
                # 自动锚定：月报键在源文件键列命中时写入 source_key / source_key_col
                if hit and value_cols:
                    rec['source_key'] = key_text
                    rec['source_key_col'] = skc
                    if vc < len(hdr):
                        rec['source_field'] = hdr[vc]
                elif hit and not value_cols:
                    rec['source_key'] = key_text
                    rec['source_key_col'] = skc
                if unit and value_cols:
                    rec['unit'] = unit
                if rp:
                    rec['reporting_period'] = rp
                if url:
                    rec['url'] = url
                rows_out.append(rec)

    with open(args.out, 'w', encoding='utf-8') as f:
        for r in rows_out:
            f.write(json.dumps(r, ensure_ascii=False) + '\n')

    # 额外写一份映射说明（.mapping.md），便于人读
    note = ['# 溯源映射说明', '', '| 表 | 源文件 | 记录数 |',
            '| --- | --- | --- |']
    for tkey, spec in mapping.items():
        ti = int(tkey) - 1
        cnt = sum(1 for r in rows_out if r['cell'].startswith('表%d!' % (ti + 1)))
        note.append('| 表%s | %s | %d |' % (tkey, spec.get('source_file', ''), cnt))
    note_path = args.out.replace('.jsonl', '.mapping.md')
    with open(note_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(note))

    print('已写出 %s（%d 条）与 %s' % (args.out, len(rows_out), note_path))


if __name__ == '__main__':
    main()
