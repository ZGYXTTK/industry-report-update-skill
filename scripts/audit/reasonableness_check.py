# -*- coding: utf-8 -*-
"""
reasonableness_check.py —— 数值合理性门禁（P0：老门禁查"自洽"，本门禁查"对不对"）

两件事，全部要求落在 --roster-note（变更摘要.md）里有交代，否则判硬伤：

① 环比异常波动：新旧月报同序号表、同业务键、同列的数值单元格，
   |变动率| > --jump-threshold（默认 0.5，即 ±50%）的，必须在变更摘要中说明原因；
② 名单变化：同键表的新增/退出键（如新增/移出赛道的公司），
   必须在变更摘要的「新增/移除标的」部分逐条点名。

设计说明：
  - 只做"找异常+要交代"，不替业务判断对错——但异常没有交代就不许过；
  - 键归一化与 diff_empty.py 对齐（半角化/去括号/去公司后缀），降低同名异写噪声；
  - 数值列按表头名对齐（容忍列序微调）；同名列取第一个。

用法:
    python reasonableness_check.py 旧月报.docx 新月报.docx \
        [--key-col-name 公司简称] [--jump-threshold 0.5] \
        --roster-note 变更摘要.md [--out 合理性校验报告.md]
退出码：0=无未交代异常；1=存在未交代异常；2=参数错误。
"""
import argparse
import re
import sys

from docx import Document

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

_NUM_RE = re.compile(r'[-+]?\d[\d,]*\.?\d*')
_AUTO_KEY_NAMES = ('公司简称', '融资方', '公司名称', '标的公司名称', '企业名称', '证券简称',
                   '股票代码', '购买方名称', '购买方', '项目')
# 表头/列里出现这些词时不做环比（序号、年份类无环比意义）
_SKIP_COL_HINTS = ('序号', '编号', '年份', '代码', '日期')


def _norm_text(v):
    if v is None:
        return ''
    out = []
    for ch in str(v):
        code = ord(ch)
        if 0xFF01 <= code <= 0xFF5E:
            ch = chr(code - 0xFEE0)
        elif code == 0x3000:
            ch = ' '
        out.append(ch)
    return re.sub(r'\s+', '', ''.join(out))


def _norm_key(v):
    t = _norm_text(v)
    t = re.sub(r'[（(][^）)]*[）)]', '', t)
    t = re.sub(r'(股份有限公司|有限责任公司|有限公司|控股集团|集团)', '', t)
    return t.strip()


# 只有这些章节里的点名才算"交代"（防把异常挪到别处消音）
_REPORT_SECTIONS = ('异常波动', '新增', '移除', '标的', '名单', '变化')


def _reportable_text(path):
    """按 Markdown 章节（# 标题）切分，只保留关键章节的正文，用于交代检查。"""
    sections = []  # (heading, body)
    cur = None
    with open(path, encoding='utf-8') as f:
        for line in f:
            s = line.strip()
            if s.startswith('#'):
                cur = s
                sections.append([cur, ''])
            elif cur is not None and s != '':
                sections[-1][1] += ' ' + s
    kept = ' '.join(body for head, body in sections
                    if any(k in _norm_text(head) for k in _REPORT_SECTIONS))
    return _norm_text(kept)


def _parse_num(v):
    """提取数值并按单位/百分比折算，返回可比较的 float；无法解析返回 None。
    支持 "12.3%"→0.123、"1.2亿元"→1.2e8、"150.80元"→150.80、"5万"→5e4。"""
    if not isinstance(v, str):
        return None
    t = v.strip().replace(',', '').replace('，', '')
    if t in ('', '-', '—', '--', '－'):
        return None
    pct = 1.0
    if t.endswith('%'):
        pct = 0.01
        t = t[:-1]
    scale = 1.0
    for u, f in (('万亿元', 1e12), ('万亿', 1e12), ('亿元', 1e8), ('亿', 1e8),
                 ('万元', 1e4), ('万', 1e4), ('元', 1.0)):
        if u in t:
            scale = f
            t = t.replace(u, '')
            break
    m = _NUM_RE.match(t)
    if not m:
        return None
    try:
        return float(m.group(0).replace(',', '')) * scale * pct
    except ValueError:
        return None


def _resolve_key_col(header, key_col_name):
    names = [key_col_name] if key_col_name else list(_AUTO_KEY_NAMES)
    for name in names:
        if not name:
            continue
        for ci, h in enumerate(header):
            if _norm_text(name) and _norm_text(name) in _norm_text(h):
                return ci
    return None


def _table_maps(table, key_col_name):
    """返回 (header, {norm_key: (原始键, [cells])})，无键列表返回 (header, None)。"""
    if not table.rows:
        return [], None
    header = [c.text.strip() for c in table.rows[0].cells]
    kc = _resolve_key_col(header, key_col_name)
    if kc is None:
        return header, None
    rows = {}
    for r in table.rows[1:]:
        cells = [c.text.strip() for c in r.cells]
        if kc >= len(cells):
            continue
        raw = cells[kc]
        nk = _norm_key(raw)
        if nk:
            rows[nk] = (raw, cells)
    return header, rows


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('old_docx')
    ap.add_argument('new_docx')
    ap.add_argument('--key-col-name', default=None)
    ap.add_argument('--jump-threshold', type=float, default=0.5, help='环比异常阈值（默认 0.5=±50%）')
    ap.add_argument('--roster-note', required=True, help='变更摘要.md 路径（异常与名单须在此交代）')
    ap.add_argument('--out', default=None)
    args = ap.parse_args()

    import os
    for p in (args.old_docx, args.new_docx, args.roster_note):
        if not os.path.exists(p):
            print('❌ 找不到 %s' % p)
            sys.exit(2)

    old_doc, new_doc = Document(args.old_docx), Document(args.new_docx)
    note_text = _reportable_text(args.roster_note)  # 只认"异常波动/新增/移除/标的/名单/变化"章节

    anomalies = []   # (表, 键, 列, 旧值, 新值, 变动率)
    added_all = []   # (表, 键)
    removed_all = []
    checked_tables = 0

    n_tables = min(len(old_doc.tables), len(new_doc.tables))
    for ti in range(n_tables):
        header_o, old_rows = _table_maps(old_doc.tables[ti], args.key_col_name)
        header_n, new_rows = _table_maps(new_doc.tables[ti], args.key_col_name)
        if old_rows is None or new_rows is None:
            continue
        checked_tables += 1
        tname = '表%d' % (ti + 1)
        for nk in sorted(set(new_rows) - set(old_rows)):
            added_all.append((tname, new_rows[nk][0]))
        for nk in sorted(set(old_rows) - set(new_rows)):
            removed_all.append((tname, old_rows[nk][0]))
        # 环比：列按表头名对齐
        col_pairs = []
        for ci, h in enumerate(header_o):
            if any(s in h for s in _SKIP_COL_HINTS):
                continue
            for cj, h2 in enumerate(header_n):
                if _norm_text(h) == _norm_text(h2) and h:
                    col_pairs.append((ci, cj, h))
                    break
        for nk in set(old_rows) & set(new_rows):
            _, cells_o = old_rows[nk]
            _, cells_n = new_rows[nk]
            for ci, cj, col_name in col_pairs:
                if ci >= len(cells_o) or cj >= len(cells_n):
                    continue
                vo, vn = _parse_num(cells_o[ci]), _parse_num(cells_n[cj])
                if vo is None or vn is None or vo == 0:
                    continue
                ratio = (vn - vo) / abs(vo)
                if abs(ratio) > args.jump_threshold:
                    anomalies.append((tname, old_rows[nk][0], col_name, vo, vn, ratio))

    # 交代检查：每个异常键、每个新增/移除键，都要在变更摘要里被点名
    unexplained = []
    for tname, key, col, vo, vn, ratio in anomalies:
        if _norm_key(key) not in note_text and _norm_text(key) not in note_text:
            unexplained.append(('异常波动未说明', '%s · %s · %s：%.4g → %.4g（%+.0f%%）' % (tname, key, col, vo, vn, ratio * 100)))
    for tname, key in added_all:
        if _norm_key(key) not in note_text and _norm_text(key) not in note_text:
            unexplained.append(('新增标的未点名', '%s · %s' % (tname, key)))
    for tname, key in removed_all:
        if _norm_key(key) not in note_text and _norm_text(key) not in note_text:
            unexplained.append(('移除标的未点名', '%s · %s' % (tname, key)))

    lines = [
        '# 数值合理性校验报告（reasonableness_check.py）',
        '',
        '- 参与比对表数：%d（新旧各 %d/%d 张）' % (checked_tables, len(old_doc.tables), len(new_doc.tables)),
        '- 环比异常阈值：±%.0f%%；命中 %d 处' % (args.jump_threshold * 100, len(anomalies)),
        '- 名单变化：新增 %d / 移除 %d' % (len(added_all), len(removed_all)),
        '- 交代依据：%s' % args.roster_note,
        '',
        '## ① 环比异常波动（需在变更摘要说明）',
        '',
        '| 表 | 键 | 列 | 上期 | 本期 | 变动 |',
        '| --- | --- | --- | --- | --- | --- |',
    ]
    for tname, key, col, vo, vn, ratio in anomalies:
        lines.append('| %s | %s | %s | %.6g | %.6g | %+.0f%% |' % (tname, key, col, vo, vn, ratio * 100))
    if not anomalies:
        lines.append('| ➖ | 无异常波动 | | | | |')
    lines += ['', '## ② 名单变化（需在变更摘要点名）', '']
    lines.append('**新增**：' + ('、'.join('%s·%s' % x for x in added_all) if added_all else '无'))
    lines.append('')
    lines.append('**移除**：' + ('、'.join('%s·%s' % x for x in removed_all) if removed_all else '无'))
    lines += ['', '## ③ 未交代清单（硬伤）', '']
    if unexplained:
        for kind, desc in unexplained:
            lines.append('- ❌ **%s**：%s' % (kind, desc))
        lines += ['', '结论：❌ %d 项异常/名单变化未在变更摘要中交代，逐条补充说明后重跑。' % len(unexplained)]
    else:
        lines.append('结论：✅ 全部异常波动与名单变化均已在变更摘要中交代。')

    report = '\n'.join(lines)
    if args.out:
        with open(args.out, 'w', encoding='utf-8') as f:
            f.write(report)
    print(report[:5000])
    print('\n===== reasonableness 汇总：波动%d 新增%d 移除%d 未交代%d =====' % (
        len(anomalies), len(added_all), len(removed_all), len(unexplained)))
    sys.exit(1 if unexplained else 0)


if __name__ == '__main__':
    main()
