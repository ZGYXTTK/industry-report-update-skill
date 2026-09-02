# -*- coding: utf-8 -*-
"""
cross_month_dedup.py —— 跨月事件去重门禁（P1-2，SKILL.md 门禁速查第 10 道）

背景（2026-08 期实测，两起阻断级）：
  1. 行业月度合集会把上月事件转载进本月（如某企业 X/23 融资轮被当成 M+1 月事件收录）；
  2. 某上市公司 M/31 公告收购 M+1 月合集转载后险些混入 M+1 月并购表。
  人工通读上月报告比对既慢又不可靠，故固化为门禁。

用法：
  # 方式一：候选为 CSV/JSONL（列含公司名，可选日期列）：
  python scripts/audit/cross_month_dedup.py --old 上月月报.docx \
      --candidates 本期候选.csv --out 去重报告.md
  # 方式二：候选以逗号分隔名单直接传入：
  python scripts/audit/cross_month_dedup.py --old 上月月报.docx \
      --names "公司A,公司B,公司C" --old-period "2026-07"

判定规则：
  - 候选名与上月报告实体名（表格业务键列 + 「涉及标的：」段落 + 含公司特征的单元格）
    按归一化键（半角化/去括号内容/去公司后缀，与 diff_empty._key 对齐）比对；
  - 命中 = 疑似跨月重复（硬信号）：必须人工核对公告日——公告日在上月 → 剔除或
    在变更摘要「逐条点名」并说明沿用/续轮口径；无法核对 → 剔除；
  - 退出码：存在疑似重复 → 1（阻断，直至逐条处置）。

退出码 0 仅表示「无名字级命中」——转载合集若改写公司名仍可能漏网，
终审以盲审自检清单第 1 条为准。
"""
import argparse
import importlib.util
import io
import os
import re
import sys

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

HERE = os.path.dirname(os.path.abspath(__file__))


def _load_diff_empty():
    spec = importlib.util.spec_from_file_location('diff_empty', os.path.join(HERE, 'diff_empty.py'))
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


def _norm(v):
    """与 diff_empty._key 一致的归一化（委托加载，保持单一实现）。"""
    return _de._key(v)


def _is_companyish(text):
    """粗略判定「像公司/机构名」：过滤表头词与纯数字序号。"""
    t = (text or '').strip().replace('\n', '')
    if not (2 <= len(t) <= 30):
        return False
    if any(s in t for s in ('序号', '简称', '合计', '总计', '来源', '备注', '日期', '金额',
                            '单位', '指标', '项目', '业务板块', '股票代码', '证券代码')):
        return False
    if re.match(r'^[\d.。%\s]+$', t):
        return False
    return True


def _extract_old_entities(docx_path):
    """从上月报告提取实体名集合：表格业务键列 + 「涉及标的：」段落 + 含公司特征单元格。"""
    from docx import Document
    doc = Document(docx_path)
    entities = {}

    def add(name, where):
        name = str(name).replace('\n', '').replace('\r', '').replace(' ', '')
        n = _norm(name)
        if n and _is_companyish(name):
            entities.setdefault(n, (name.strip(), where))

    for ti, tbl in enumerate(doc.tables):
        kc = 0
        try:
            kc = de._resolve_key_idx(tbl) if hasattr(de, '_resolve_key_idx') else 0
        except Exception:
            kc = 0
        for row in tbl.rows[1:]:
            seen = set()
            for ci, c in enumerate(row.cells):
                t = c.text.strip().replace('\n', '').replace(' ', '')
                if t and t not in seen:
                    seen.add(t)
                    if ci == kc or '公司' in t or '科技' in t or '航空' in t or '航天' in t \
                            or '智能' in t or '空间' in t or '星' in t or '机器人' in t:
                        add(t, '表%d' % (ti + 1))
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith('涉及标的：'):
            add(t[len('涉及标的：'):].strip(), '正文涉及标的')
    return entities


def _load_candidates(path):
    """候选：CSV（列含公司名称/公司/企业名称/融资方/名称）或 JSONL（同名字段）。"""
    import csv as _csv
    cands = []
    if path.endswith('.jsonl'):
        import json
        for line in open(path, encoding='utf-8'):
            line = line.strip()
            if not line:
                continue
            r = json.loads(line)
            name = next((r.get(k) for k in ('公司名称', '公司', '企业名称', '融资方', '公司简称',
                                            '证券简称', 'name') if r.get(k)), None)
            date = next((r.get(k) for k in ('获投日期', '日期', '披露日期', 'date', '公告日') if r.get(k)), '')
            if name:
                cands.append((str(name), str(date)))
        return cands
    for enc in ('utf-8-sig', 'utf-8', 'gbk'):
        try:
            with open(path, encoding=enc, newline='') as f:
                rows = [r for r in _csv.reader(f) if any((c or '').strip() for c in r)]
            break
        except (UnicodeDecodeError, UnicodeError):
            continue
    else:
        return cands
    if not rows:
        return cands
    header = [h.strip() for h in rows[0]]
    name_ci = next((ci for ci, h in enumerate(header)
                    if any(k in h for k in ('公司名称', '公司简称', '企业名称', '融资方', '名称', '购买方'))), 0)
    date_ci = next((ci for ci, h in enumerate(header) if '日期' in h), None)
    for r in rows[1:]:
        name = r[name_ci].strip() if name_ci < len(r) else ''
        date = r[date_ci].strip() if date_ci is not None and date_ci < len(r) else ''
        if name:
            cands.append((name, date))
    return cands


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--old', required=True, help='上月月报 docx')
    ap.add_argument('--candidates', default=None, help='本期候选 CSV/JSONL（含公司名列）')
    ap.add_argument('--names', default=None, help='逗号分隔候选名单（与 --candidates 二选一）')
    ap.add_argument('--old-period', default='', help='上月报告期（如 2026-07），仅用于报告展示')
    ap.add_argument('--out', default=None)
    args = ap.parse_args()

    global _de
    if _de is None:
        _de = _load_diff_empty()

    if not os.path.exists(args.old):
        print('❌ 找不到上月月报：%s' % args.old)
        return 2

    cands = []
    if args.candidates:
        cands = _load_candidates(args.candidates)
    if args.names:
        cands += [(n.strip(), '') for n in args.names.split(',') if n.strip()]
    if not cands:
        print('❌ 无候选（--candidates 或 --names 至少提供一个）')
        return 2

    entities = _extract_old_entities(args.old)

    hits = []
    for name, date in cands:
        nk = _norm(name)
        if not nk:
            continue
        if nk in entities:
            old_name, where = entities[nk]
            hits.append((name, date, old_name, where, '全名命中'))
            continue
        # 双向子串匹配（简称 vs 全称：候选「A 公司」⊂ 上月「北京 A 公司科技有限公司」）
        if len(nk) >= 4:
            for ek, (old_name, where) in entities.items():
                if len(ek) >= 4 and (nk in ek or ek in nk):
                    hits.append((name, date, old_name, where, '子串命中'))
                    break

    lines = ['# 跨月去重报告（cross_month_dedup.py · 门禁第 10 道）', '',
             '| 指标 | 值 |', '| --- | --- |',
             '| 上月报告 | %s |' % os.path.basename(args.old),
             '| 上月期号 | %s |' % (args.old_period or '未指定'),
             '| 本期候选数 | %d |' % len(cands),
             '| 上月实体名数 | %d |' % len(entities),
             '| 疑似跨月重复 | %d |' % len(hits),
             '| 结论 | %s |' % ('❌ 存在疑似重复——逐条核对公告日：上月已披露 → 剔除或点名说明；无法核对 → 剔除'
                              if hits else '✅ 无名字级命中（转载合集改名仍可能漏网，终审以盲审清单第 1 条为准）'), '']
    if hits:
        lines.append('## 疑似跨月重复（必处置）')
        for name, date, old_name, where, how in hits:
            lines.append('- ❌ 候选「%s」%s ↔ 上月已出现「%s」（%s，%s）' % (name, ('[%s]' % date) if date else '', old_name, where, how))
        lines.append('')
        lines.append('> 处置：①公告日在上月 → 从本期表剔除，并在变更摘要注明；'
                     '> ②同轮次新增交割/补充披露 → 保留并在轮次列标注「（续）」，变更摘要逐条点名；'
                     '> ③确属本月新事件（名字相近但主体不同）→ 在变更摘要写明区分依据。')
    report = '\n'.join(lines)
    if args.out:
        with open(args.out, 'w', encoding='utf-8') as f:
            f.write(report)
    print(report)
    return 1 if hits else 0


_de = None

if __name__ == '__main__':
    sys.exit(main())
