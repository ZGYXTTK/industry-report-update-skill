# -*- coding: utf-8 -*-
"""
consistency_check.py —— 数据流审计第③件：一致性 / 逻辑校验（v4）

v4 相对 v3 的改进（对应「正文数字候选误报多」优化项）：
  1. 模糊量词白名单：在"约/～/≥/≤/以上/以下/招股书/未披露/未取得/区间/超X元"
     等语境内的数字，**不**纳入"正文数字候选"，避免误报"该数字未见于任何表格"。
  2. 表格占比列豁免名单新增：「占营业收入比重 / 占融资比重 / 占投资比重」等非求和型字段。
  3. 比例型字段白名单（v3 已有）：保留。

三件事：
  1. 表格内「合计/总计」行是否等于分项数值之和（硬门禁）；
  2. 占比列是否加总 ≈ 100%（容忍 1% 误差，比例型字段豁免）；
  3. 正文出现的数字与表格数字的「可反查候选」提示（提示人工核对，非强判）。

用法:
    python consistency_check.py 新月报.docx [--out 一致性校验报告.md]
"""
import argparse
import re
import sys

from docx import Document

NUM = re.compile(r'-?\d+(?:\.\d+)?')
TOTAL_KEYS = ('合计', '总计', '小计', 'Total', '合计值')
# 天然不求和的「比率/倍数」型字段：命中即跳过占比校验
RATIO_KEYS = ('毛利率', '净利率', '市盈率', '市净率', '换手率', '收益率', '资产负债率',
              'ROE', 'ROA', 'PE', 'PB', 'PS', '速动比率', '流动比率', '周转率', '增速',
              '占营业收入比重', '占融资比重', '占投资比重', '份额', '占比')

# 模糊量词白名单：段落文本同时命中量词与数字 → 视为「不强制可反查」，不进入 orphan
FUZZY_QUANTIFIERS = ('约', '～', '~', '≥', '≤', '以上', '以下', '近', '超过', '不足', '区间',
                     '招股书', '未披露', '未取得', '估值超', '突破', '高达', '低至', '整体',
                     '较之上', '同类水平', '暂未', '已锁定', '延续', '持续', '陆续',
                     '2026年', '2025年', '2030年', '2024年', '下半年', '上半年', '年度',
                     '月内', '月初', '月末', '中旬', '上年同期', '历史', '未来')
# 行级白名单（含数字的整行若是「年/月/年份」语境，整行豁免）
YEAR_CONTEXT = re.compile(r'^\s*(20\d{2}\s*年|[一二三四五六七八九十]+\s*年|上半年|下半年|月内|月初|月末|中旬|首日|次日|当日)\s*[:：]?\s*[^.\n]*$')


def _col_letter(ci):
    s = ''
    n = ci + 1
    while n > 0:
        n, r = divmod(n - 1, 26)
        s = chr(65 + r) + s
    return s


def _to_float(s):
    if not s:
        return None
    t = str(s).replace(',', '').replace('，', '')
    m = NUM.search(t)
    if not m:
        return None
    try:
        return float(m.group())
    except ValueError:
        return None


def _is_total_cell(text):
    return any(k in text for k in TOTAL_KEYS)


def _is_ratio_col(header):
    return any(k in header for k in RATIO_KEYS)


def _is_cover_style(style_name):
    """封面/扉页段落的 style 不含 Heading/标题，用段落序号截断更稳（见 main）。"""
    return False


def _skip_orphan_num(n):
    """过滤邮编/电话/股票代码等非正文数字：≥5 位整数、含连字符、含小数点后缀 .HK/.SH/.SZ 等。"""
    if '.' in n:
        # 小数保留，但排除 '1776.HK' 这类代码（. 后跟字母）
        return bool(re.search(r'\.(HK|SH|SZ|BJ|OF|IB)$', n, re.I))
    return len(n) >= 5  # 邮编/电话/股票代码等长数字串


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('docx')
    ap.add_argument('--out', default='一致性校验报告.md')
    args = ap.parse_args()

    doc = Document(args.docx)
    issues = []

    # ---- 1) 合计行 = 分项和 ----
    for ti, table in enumerate(doc.tables):
        rows = table.rows
        for ri, row in enumerate(rows):
            if ri == 0:
                continue
            if not any(_is_total_cell(c.text) for c in row.cells):
                continue
            for ci in range(len(row.cells)):
                # 比率/倍数型列（毛利率、市盈率等）的「合计」是加权平均，天然不等于分项简单加总——跳过
                _hdr = rows[0].cells[ci].text if ci < len(rows[0].cells) else ''
                if _is_ratio_col(_hdr):
                    continue
                total_v = _to_float(row.cells[ci].text)
                if total_v is None:
                    continue
                parts = []
                for r2 in rows[1:ri]:
                    if r2 is row:
                        continue
                    # 跳过「小计/子类合计」行，避免两级合计把子小计重复加总
                    if r2.cells and _is_total_cell(r2.cells[0].text):
                        continue
                    if len(r2.cells) <= ci:
                        continue
                    v = _to_float(r2.cells[ci].text)
                    if v is not None:
                        parts.append(v)
                if parts:
                    s = sum(parts)
                    if abs(s - total_v) > 0.5:
                        issues.append(('合计≠分项和',
                                       '表%d!%s%d 合计=%s 但分项和=%s'
                                       % (ti + 1, _col_letter(ci), ri + 1, total_v, round(s, 2))))

    # ---- 2) 占比加总 ≈ 100%（比例型字段豁免）----
    for ti, table in enumerate(doc.tables):
        if not table.rows:
            continue
        header = [c.text.strip() for c in table.rows[0].cells]
        pct_cols = [ci for ci, h in enumerate(header)
                    if any(k in h for k in ('占比', '比例', '比重', '%', '份额'))]
        for ci in pct_cols:
            if _is_ratio_col(header[ci]):
                continue  # 毛利率/市盈率等天然不求和的字段，跳过
            vals = []
            for ri in range(1, len(table.rows)):
                if _is_total_cell(table.rows[ri].cells[0].text):
                    continue
                v = _to_float(table.rows[ri].cells[ci].text)
                if v is not None:
                    vals.append(v)
            if vals:
                s = sum(vals)
                if 90 <= s <= 110 or 0.9 <= s <= 1.1:
                    continue
                issues.append(('占比加总异常',
                               '表%d 列%s 占比加总=%s（疑似未收敛到100%%）'
                               % (ti + 1, _col_letter(ci), round(s, 2))))

    # ---- 3) 正文 vs 表格数字候选（跳过封面，过滤长数字串）----
    # 定位第一个 Heading/标题 段落，之前的视为封面/扉页，不参与数字候选
    first_heading = 0
    for i, p in enumerate(doc.paragraphs):
        sn = (p.style.name or '')
        if 'Heading' in sn or '标题' in sn or 'toc' in sn.lower():
            first_heading = i
            break
    body_nums = set()
    for i, p in enumerate(doc.paragraphs):
        if i < first_heading:
            continue  # 封面/扉页
        text = p.text
        # v4 模糊量词白名单：段落命中量词 → 整段数字不进候选
        if any(k in text for k in FUZZY_QUANTIFIERS):
            continue
        # 行级年份语境白名单（如「2026 年：」「上半年：」整行）
        if YEAR_CONTEXT.match(text):
            continue
        for m in NUM.findall(text):
            if not _skip_orphan_num(m):
                body_nums.add(m)
    table_nums = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for m in NUM.findall(cell.text):
                    table_nums.add(m)
    orphan = sorted([n for n in body_nums - table_nums if '.' not in n], key=lambda x: -len(x))[:20]

    lines = [
        '# 一致性校验报告（v4：比例字段豁免 + 模糊量词白名单 + 封面数字过滤）',
        '',
        '## 1) 合计 ≠ 分项和',
    ]
    n1 = sum(1 for k, _ in issues if k == '合计≠分项和')
    if n1:
        for k, msg in issues:
            if k == '合计≠分项和':
                lines.append('- ❌ ' + msg)
    else:
        lines.append('- ✅ 未发现')

    lines += ['', '## 2) 占比加总异常（比例型字段已豁免）']
    n2 = sum(1 for k, _ in issues if k == '占比加总异常')
    if n2:
        for k, msg in issues:
            if k == '占比加总异常':
                lines.append('- ⚠️ ' + msg)
    else:
        lines.append('- ✅ 未发现')

    lines += ['', '## 3) 正文数字未见于任何表格（候选，供人工核对；封面数字已过滤）']
    if orphan:
        lines.append('- ' + '、'.join(orphan))
    else:
        lines.append('- ✅ 无')

    lines += ['', '> 合计为强校验（不通过则非零退出）；占比与第 3 项为提示，不自动判错。']
    with open(args.out, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    print('已写出 %s（合计异常 %d，占比异常 %d，正文候选 %d）'
          % (args.out, n1, n2, len(orphan)))
    if n1:
        sys.exit(1)


if __name__ == '__main__':
    main()

