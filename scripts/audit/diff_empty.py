# -*- coding: utf-8 -*-
"""
diff_empty.py —— 数据流审计第②件：空值 diff（防漏填，键列对齐版，v3）

v3 相对 v2 的改进（对应「键列用序号导致跨月误报」优化项）：
  1. 新增 --key-col-name：按「表头名称」定位业务主键列（如 "公司简称"/"标的公司名称"/
     "融资方"），优先于 --key-col（序号下标）。用业务键对齐，同一家跨月丢字段才会被揪出，
     「换了一家公司」不再误报；
  2. 键值归一化：对齐前 strip 空白，降低「A 公司」vs「A 公司（ABC）」类噪声；
  3. 报告区分「真缺陷（同键→空值）」与「实体变更（替换/增删行）」，后者单独列示。

用法:
    python diff_empty.py 旧月报.docx 新月报.docx [--key-col 0] [--key-col-name 公司简称] [--out 空值对比报告.md]
"""
import argparse
import difflib
import re
import sys

from docx import Document


def _col_letter(ci):
    s = ''
    n = ci + 1
    while n > 0:
        n, r = divmod(n - 1, 26)
        s = chr(65 + r) + s
    return s


def _cell_ok(text):
    t = (text or '').strip()
    return t not in ('', '-', '—', '--', '－')


def _key(text):
    """归一化键值：全角转半角、去空白、去括号及括号内容、去公司后缀，降低同名异写噪声。"""
    t = (text or '')
    # 全角转半角
    out = []
    for ch in t:
        code = ord(ch)
        if 0xFF01 <= code <= 0xFF5E:
            ch = chr(code - 0xFEE0)
        elif code == 0x3000:
            ch = ' '
        out.append(ch)
    t = ''.join(out)
    t = t.strip()
    import re as _re
    t = _re.sub(r'[（(][^）)]*[）)]', '', t)          # 去括号及内容（"A 公司（ABC）"→"A 公司"）
    t = _re.sub(r'(股份有限公司|有限公司|集团|有限责任公司|控股集团)', '', t)  # 去公司后缀
    t = t.strip()
    return t


# 常见业务键列名（当 --key-col-name 未传时按此顺序自动识别，避免默认序号列假阳性）
# 顺序：优先公司简称类、其次融资方、再次股票代码/购买方等
_AUTO_KEY_NAMES = ('公司简称', '融资方', '公司名称', '标的公司名称', '企业名称', '证券简称',
                   '股票代码', '购买方名称', '购买方', '融资方主营业务', '项目')


def _hdr_norm(s):
    """表头归一化：去所有空白（含单元格内换行 ␤），'公司\n简称'→'公司简称'。"""
    return re.sub(r'\s+', '', (s or ''))


def _resolve_key_idx(table, key_col, key_col_name):
    """优先按表头名称定位键列；未命中时按常见业务键列名**精确匹配**自动识别；最后才回退 --key-col 下标。

    2026-09 修复：此前 run.py 强制传 --key-col-name 公司简称，行业动态表（键列=标的公司名称）
    匹配不到「公司简称」，直接退回 --key-col=0（序号），造成跨月事件表大面积「序号对齐误报」。
    现改为：① key_col_name 未命中 → 继续 _AUTO_KEY_NAMES 按本表表头自动识别；② 匹配用
    「去空白后精确相等」而非子串包含——否则「融资方」会误配「融资方式」、「购买方」误配「购买方主营业务」、
    「公司简称」因表头换行（公司␤简称）匹配不到而退回序号。
    """
    if table.rows:
        header = [_hdr_norm(c.text) for c in table.rows[0].cells]
        if key_col_name:
            kn = _hdr_norm(key_col_name)
            for ci, h in enumerate(header):
                if h == kn:
                    return ci
        # 关键修复：key_col_name 未命中时，继续按表头精确匹配自动识别业务键（勿退回序号）
        for name in _AUTO_KEY_NAMES:
            nn = _hdr_norm(name)
            for ci, h in enumerate(header):
                if h == nn:
                    return ci
    return key_col


def _row_keys(rows, key_idx):
    keys = []
    for r in rows:
        if key_idx < len(r.cells):
            keys.append(_key(r.cells[key_idx].text))
        else:
            keys.append('')
    return keys


def _header_sig(table):
    """表头签名：第一行非空 cell 的归一化元组（用于跨月按表头对齐，而非按索引）。"""
    if not table.rows:
        return ()
    return tuple(_key(c.text) for c in table.rows[0].cells if c.text.strip())


def _match_tables(old_doc, new_doc):
    """按表头签名对齐新旧表，返回 [(old_table, new_table, 匹配方式)]；签名对不上的单独列出。"""
    old_tabs = list(old_doc.tables)
    new_tabs = list(new_doc.tables)
    old_sigs = [_header_sig(t) for t in old_tabs]
    new_sigs = [_header_sig(t) for t in new_tabs]
    used_new = set()
    pairs = []
    unmatched = []
    for oi, osig in enumerate(old_sigs):
        # 优先精确签名匹配，退化为序号匹配
        hit = None
        for ni, nsig in enumerate(new_sigs):
            if ni not in used_new and nsig == osig:
                hit = ni
                break
        if hit is None and oi < len(new_tabs) and oi not in used_new:
            hit = oi  # 退化：同序号
        if hit is not None:
            used_new.add(hit)
            pairs.append((oi, hit))
        else:
            unmatched.append(('旧表%d 无匹配' % (oi + 1), osig))
    for ni in range(len(new_tabs)):
        if ni not in used_new:
            unmatched.append(('新表%d 为新增表' % (ni + 1), new_sigs[ni]))
    return pairs, unmatched


def _compare_row(ti, old_row, new_row, key_text, defects):
    """逐列比较两行，凡「旧有值→新空值」记为缺陷（必补）。"""
    nc = min(len(old_row.cells), len(new_row.cells))
    for ci in range(nc):
        ov = old_row.cells[ci].text
        nv = new_row.cells[ci].text
        if _cell_ok(ov) and not _cell_ok(nv):
            loc = '表%d[%s]%s' % (ti + 1, key_text, _col_letter(ci))
            defects.append((loc, ov.strip(), nv.strip()))


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('old_docx')
    ap.add_argument('new_docx')
    ap.add_argument('--key-col', type=int, default=0, help='键列下标（0=第一列），--key-col-name 优先')
    ap.add_argument('--key-col-name', default='',
                    help='按表头名称定位业务键列（如 "公司简称"/"标的公司名称"/"融资方"）。留空=按 _AUTO_KEY_NAMES 表头精确匹配自动识别（推荐）；仅当某表表头无业务键列时才显式指定')
    ap.add_argument('--out', default='空值对比报告.md')
    args = ap.parse_args()

    old_doc = Document(args.old_docx)
    new_doc = Document(args.new_docx)

    defects = []
    row_notes = []
    key_used = []
    pairs, unmatched = _match_tables(old_doc, new_doc)
    for oi, ni in pairs:
        ot, nt = old_doc.tables[oi], new_doc.tables[ni]
        key_idx = _resolve_key_idx(nt, args.key_col, args.key_col_name)
        if nt.rows and key_idx < len(nt.rows[0].cells):
            key_name = nt.rows[0].cells[key_idx].text.strip()
        else:
            key_name = '列%d' % key_idx
        key_used.append('旧表%d↔新表%d 键列=%s' % (oi + 1, ni + 1, key_name))
        old_data = ot.rows[1:] if len(ot.rows) > 1 else []
        new_data = nt.rows[1:] if len(nt.rows) > 1 else []
        old_keys = _row_keys(old_data, key_idx)
        new_keys = _row_keys(new_data, key_idx)
        sm = difflib.SequenceMatcher(None, old_keys, new_keys)
        for tag, i1, i2, j1, j2 in sm.get_opcodes():
            if tag == 'equal':
                for a, b in zip(range(i1, i2), range(j1, j2)):
                    _compare_row(oi, old_data[a], new_data[b], old_keys[a], defects)
            elif tag == 'delete':
                row_notes.append('旧表%d→新表%d：删除行（旧）%s'
                                 % (oi + 1, ni + 1, '、'.join(old_keys[i1:i2]) or '（键列为空）'))
            elif tag == 'insert':
                row_notes.append('旧表%d→新表%d：新增行（新）%s'
                                 % (oi + 1, ni + 1, '、'.join(new_keys[j1:j2]) or '（键列为空）'))
            elif tag == 'replace':
                row_notes.append('旧表%d→新表%d：替换行（实体变更，不做逐格比较）旧[%s] → 新[%s]'
                                 % (oi + 1, ni + 1, '、'.join(old_keys[i1:i2]), '、'.join(new_keys[j1:j2])))

    if unmatched:
        for msg, sig in unmatched:
            row_notes.append('⚠️ 表头签名未匹配：%s（表头=%s）' % (msg, '|'.join(sig)[:60]))

    lines = [
        '# 空值对比报告（业务键对齐；旧有值 → 新空值 = 缺陷）',
        '',
        '## 键列定位',
    ]
    for k in key_used:
        lines.append('- ' + k)
    lines += ['', '## 缺陷清单（同键→空值，必补）',
              '| # | 位置 | 旧值 | 新值 | 处置 |',
              '| --- | --- | --- | --- | --- |']
    for i, (loc, ov, nv) in enumerate(defects, 1):
        lines.append('| %d | %s | %s | %s | 必补 |' % (i, loc, ov.replace('|', '\\|'), nv))
    if not defects:
        lines.append('| — | 无 | — | — | — |')
    lines += ['', '## 实体变更说明（增删/替换行，不计为缺陷）']
    if row_notes:
        for n in row_notes:
            lines.append('- ' + n.replace('|', '\\|'))
    else:
        lines.append('- 无')
    lines += ['', '> 共 %d 处缺陷。凡「旧有值→新空值」必须补（法定代表人/PE-VC 等用 qcc 核实）。'
              % len(defects),
              '> 若「替换行」明显增多，说明键列选错（业务键被当序号对齐），请改用 --key-col-name 指定业务键。',
              '> 注意：本门禁只抓「同业务键跨月丢字段」；事件/融资类表整批换血后的**新行自身漏填**',
              '> 由 field_completeness.py（P0）负责，两者互补，不可互相替代。']

    with open(args.out, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    print('已写出 %s，共发现 %d 处缺陷，%d 条实体变更说明' % (args.out, len(defects), len(row_notes)))
    if defects:
        sys.exit(1)


if __name__ == '__main__':
    main()
