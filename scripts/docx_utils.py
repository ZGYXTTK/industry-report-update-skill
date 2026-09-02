# -*- coding: utf-8 -*-
"""
docx_utils.py —— 格式保真工具库（industry-report-update 专用）

相比 SKILL 早期内联代码，本库补齐三处工程短板：
  1. 多 run 单元格：逐 run 克隆 rPr（而非只复制首个 run，避免「整格变粗/变字号」）；
  2. 多段落单元格：逐段克隆 pPr（而非只复制首段对齐）；
  3. 合并单元格：add_row 时按网格重建 gridSpan / vMerge（避免加行后列错位）。

依赖：python-docx（本机已装）。
使用：from docx_utils import set_cell_text_keep_fmt, add_row_copy_fmt
"""
from copy import deepcopy

from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ----------------------------------------------------------------------
# 基础 XML 工具
# ----------------------------------------------------------------------
def _find(el, tag):
    return el.find(qn(tag))


def _remove(el, tag):
    old = _find(el, tag)
    if old is not None:
        el.remove(old)
    return old


def _clone_into(el, tag, src_el):
    """把 src_el 的指定子元素深拷贝插入 el（替换同名旧子元素）。"""
    src = _find(src_el, tag)
    if src is None:
        return
    _remove(el, tag)
    el.insert(0, deepcopy(src))


# ----------------------------------------------------------------------
# 段落
# ----------------------------------------------------------------------
def _all_run_elements(p):
    """返回段落下所有 <w:r> 元素，含 <w:hyperlink> 内部的 run，按文档顺序。"""
    els = []
    for child in p._element:
        if child.tag == qn('w:r'):
            els.append(child)
        elif child.tag == qn('w:hyperlink'):
            for sub in child:
                if sub.tag == qn('w:r'):
                    els.append(sub)
    return els


def _set_run_text(el, text):
    """向 <w:r> 写入文本（复用首个 w:t，清空其余）。"""
    ts = el.findall(qn('w:t'))
    if not ts:
        t = el.makeelement(qn('w:t'), {})
        t.text = text
        el.append(t)
    else:
        ts[0].text = text
        for t in ts[1:]:
            t.text = ""


def set_para_text_keep_fmt(p, new_text):
    """改段落文字但保留段落格式(pPr)与首个 run 格式(rPr)，兼容超链接内 run。"""
    runs = _all_run_elements(p)
    if runs:
        _set_run_text(runs[0], new_text)
        for el in runs[1:]:
            _set_run_text(el, "")
    else:
        p.add_run(new_text)


def _apply_rpr_overrides(run, overrides):
    """把 overrides（如 {'b': True} 或 {'sz': 21}）应用到 run 的 rPr。"""
    rpr = _find(run._element, 'w:rPr')
    if rpr is None:
        rpr = run._element.makeelement(qn('w:rPr'), {})
        run._element.insert(0, rpr)
    for k, v in overrides.items():
        tag = qn('w:%s' % k)
        el = rpr.find(tag)
        if v is True:
            if el is None:
                el = rpr.makeelement(tag, {})
                rpr.append(el)
        elif v is False or v is None:
            if el is not None:
                rpr.remove(el)
        else:
            if el is None:
                el = rpr.makeelement(tag, {})
                rpr.append(el)
            el.set(qn('w:val'), str(v))


def set_para_segments_keep_fmt(p, segments):
    """
    按多段格式回填段落（「标题加粗 run + 正文不加粗 run」这类）。
    segments: list[tuple[str, dict|None]]，dict 为可选 rPr 覆盖（如 {"b": True}）。
    首个 run 的 rPr 作为基准，其余 run 克隆基准再按需覆盖。
    """
    if not p.runs:
        p.add_run("")
    base_rpr = deepcopy(_find(p.runs[0]._element, 'w:rPr'))
    # 移除现有 run（含超链接内的 run）
    for el in _all_run_elements(p):
        el.getparent().remove(el)
    for text, override in segments:
        run = p.add_run(text)
        if base_rpr is not None:
            run._element.insert(0, deepcopy(base_rpr))
        if override:
            _apply_rpr_overrides(run, override)


# ----------------------------------------------------------------------
# 单元格
# ----------------------------------------------------------------------
def set_cell_text_keep_fmt(cell, value):
    """改单元格文字但保留 tcPr(vAlign)/pPr(jc)/rPr(字号)，兼容超链接内 run。

    v2（2026-09 实测修正）：多段落单元格此前只改首段，旧值残留于后续段落——
    是内容类门禁（一致性/空值 diff）误报的根因。现对全部段落执行 run 保留式写入：
    首段写 value，其余段落文本置空（段落/run 结构原样保留，防格式回归）。
    """
    for i, p in enumerate(cell.paragraphs):
        set_para_text_keep_fmt(p, value if i == 0 else "")


def set_cell_keep(cell, value):
    """单元格格式保真写入 v3：保留全部段落与 run 结构。

    - value 为 str：写入首段首 run，其余段落/run 文本置空（与 set_cell_text_keep_fmt 等价）；
    - value 为 list/tuple：按段落 1:1 分配（超出段落数的行并入末段、以「；」连接；
      不足的段落置空）——用于多行结构化单元格（如「注册资本/成立日期/地址」）。

    构建脚本应优先使用本函数族（set_cell_keep/fill_table），禁止 cell.text= 覆写。
    """
    lines = list(value) if isinstance(value, (list, tuple)) else [value]
    ps = cell.paragraphs
    n = len(ps)
    if len(lines) > n:
        lines = lines[:n - 1] + ["；".join(str(x) for x in lines[n - 1:])]
    for i, p in enumerate(ps):
        set_para_text_keep_fmt(p, str(lines[i]) if i < len(lines) else "")


def set_para_keep(p, text):
    """段落级格式保真写入：首 run 写文本、其余 run 文本置空（run 结构原样保留）。

    空文本且无 run 的段落原样跳过。整段重建（set_para_segments_keep_fmt 单 segment、
    p.add_run 新建）会破坏 run 序列——2026-08 期实测使 format_diff 跌至 94.1%，
    run 保留式写回后 96.6%。构建脚本改写既有段落时应用本函数。
    """
    if not text and not _all_run_elements(p):
        return p
    return set_para_text_keep_fmt(p, text)


def strip_vmerge(table):
    """删除全表所有 vMerge（纵向合并）标记，使每个单元格独立；保留 gridSpan。

    用于：复用上月模板表格重填数据前清除历史合并——续表「序号/事件」列的
    vMerge 残留会导致序号跳号、事件-标的错位（2026-08 期实测盲审阻断级根因）。
    返回删除的 vMerge 数量。
    """
    count = 0
    for tcPr in table._tbl.iter(qn('w:tcPr')):
        for vm in list(tcPr.findall(qn('w:vMerge'))):
            tcPr.remove(vm)
            count += 1
    return count


def fill_table(table, rows, start_row=1, keep_surplus=False):
    """数据行通用回填：删多余数据行（自底向上）→ 不足时按末行深拷贝克隆 → 逐格 set_cell_keep。

    rows: list[sequence]，每项为该行按列的取值（str 或 list[多段落]）。
    解决两类结构问题（2026-08 期实测）：重填后旧数据行残留（一致性门禁误报）、
    行数不足时手工克隆遗漏。返回 (删除行数, 克隆行数)。

    注意：含 gridSpan/vMerge 的复杂表请先 strip_vmerge；gridSpan 表格的 rows
    项数应与 row.cells 展开数一致。
    """
    deleted = 0
    while len(table.rows) - start_row > len(rows):
        tr = table.rows[-1]._tr
        tr.getparent().remove(tr)
        deleted += 1
    cloned = 0
    while len(table.rows) - start_row < len(rows):
        last = table.rows[-1]._tr
        last.addnext(deepcopy(last))
        cloned += 1
    for ri, vals in enumerate(rows):
        cells = table.rows[start_row + ri].cells
        if len(vals) > len(cells):
            raise ValueError('行 %d 值数 %d 超过列数 %d'
                             % (start_row + ri, len(vals), len(cells)))
        for ci, v in enumerate(vals):
            set_cell_keep(cells[ci], v)
    return deleted, cloned


def _clone_paragraphs(src_cell, dst_cell):
    """清空目标单元格段落，逐段克隆源单元格段落的 pPr + 每个 run 的 rPr。"""
    for p in list(dst_cell.paragraphs):
        p._element.getparent().remove(p._element)
    for sp in src_cell.paragraphs:
        new_p = dst_cell.add_paragraph()
        _clone_into(new_p._element, 'w:pPr', sp._element)
        for r in sp.runs:
            nr = new_p.add_run(r.text)
            _clone_into(nr._element, 'w:rPr', r._element)


def copy_cell_fmt(src_cell, dst_cell):
    """
    深拷贝单元格全套格式：tcPr（含 vAlign、gridSpan、vMerge）+ 每段 pPr + 每 run rPr。
    相比旧版（只复制首 run/首段），对多段落/多 run/合并单元格均正确。
    """
    _clone_into(dst_cell._tc, 'w:tcPr', src_cell._tc)
    _clone_paragraphs(src_cell, dst_cell)


def _row_grid(row):
    """
    返回该行按网格展开的 [(tc_element, grid_index, span), ...]。
    处理 gridSpan：一个 tc 占多列时只返回其起始网格下标，span 从 tcPr 读取。
    """
    result = []
    grid_idx = 0
    for tc in row._tr.tc_lst:
        tcPr = _find(tc, 'w:tcPr')
        span = 1
        if tcPr is not None:
            gs = _find(tcPr, 'w:gridSpan')
            if gs is not None:
                span = int(gs.get(qn('w:val'), '1'))
        result.append((tc, grid_idx, span))
        grid_idx += span
    return result


def add_row_copy_fmt(table, values, src_row_index=None):
    """
    新增数据行：按源行「逻辑列」重建单元格（含 gridSpan），再复制格式填值。

    v2 修正：`table.add_row()` 会按 grid 列数建物理 tc（每列一个），
    若把含 gridSpan 的源 tcPr 盲拷到这些物理 tc 上，会网格溢出/列错位。
    正确做法：清掉 add_row 默认建的 tc，按源行逻辑列逐个重建 tc，
    再复制对应源单元格的 tcPr（含 gridSpan）/pPr/rPr。

    values 按源行「逻辑列」顺序；src_row_index 不传时取最后一条数据行作模板。
    """
    if len(table.rows) < 2:
        raise ValueError('表格需至少包含表头 + 一条数据行作为格式模板')
    src_row = table.rows[src_row_index if src_row_index is not None else len(table.rows) - 1]
    src_grid = _row_grid(src_row)  # [(src_tc, grid_idx, span), ...]
    if len(values) != len(src_grid):
        raise ValueError('values 数量 %d 与源行逻辑列数 %d 不一致，请按逻辑列填值'
                         % (len(values), len(src_grid)))

    tr = table.add_row()
    # 清掉 add_row 默认建的物理 tc，改按源行逻辑列重建
    for tc in list(tr._tr.tc_lst):
        tr._tr.remove(tc)
    for v, (src_tc, _grid_idx, _span) in zip(values, src_grid):
        tc = OxmlElement('w:tc')
        tr._tr.append(tc)
        # 复制源单元格 tcPr（含 gridSpan / vAlign）——此时 tc 按逻辑列创建，复制才正确
        _clone_into(tc, 'w:tcPr', src_tc)
        _fill_tc_from_template(tc, src_tc, v)
    return tr


def _fill_tc_from_template(dst_tc, src_tc, value):
    """以源单元格第一段为模板：复制 pPr + 首 run rPr，回填文字到新建 run。"""
    p = OxmlElement('w:p')
    dst_tc.append(p)
    src_ps = src_tc.findall(qn('w:p'))
    if src_ps:
        _clone_into(p, 'w:pPr', src_ps[0])
    src_rs = src_ps[0].findall(qn('w:r')) if src_ps else []
    r = OxmlElement('w:r')
    p.append(r)
    if src_rs:
        _clone_into(r, 'w:rPr', src_rs[0])
    t = OxmlElement('w:t')
    t.text = str(value)
    r.append(t)
